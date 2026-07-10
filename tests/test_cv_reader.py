"""Tests for DOCX table extraction."""

from docx import Document

from src.parser.cv_reader import read_cv_document


def test_docx_reads_tables_in_document_order(tmp_path):
    path = tmp_path / "cv.docx"
    doc = Document()
    doc.add_paragraph("Professional Summary")
    doc.add_paragraph("Telecom consultant with international experience.")

    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Year"
    table.rows[0].cells[1].text = "Degree"
    table.rows[0].cells[2].text = "Institution"
    table.rows[1].cells[0].text = "2010"
    table.rows[1].cells[1].text = "B.E. Electrical Engineering"
    table.rows[1].cells[2].text = "University of Engineering"
    doc.save(path)

    text = read_cv_document(path)
    assert "Professional Summary" in text
    assert "[TABLE]" in text
    assert "B.E. Electrical Engineering" in text
    assert "University of Engineering" in text
    assert text.index("Professional Summary") < text.index("[TABLE]")
